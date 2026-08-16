#!/usr/bin/env python3
"""
Render the location index as a Word document.

Same data and the same category mapping as make_location_index.py — the codes
and labels are imported from it rather than restated, so the .sql and the .docx
can never disagree.

    python make_location_index_docx.py                     # blank address column
    python make_location_index_docx.py data/addresses.csv  # filled in

Output: Campus Location Index.docx
"""

import csv
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from make_location_index import CODE, LABEL, load_addresses

HERE = Path(__file__).parent
LOCATIONS_CSV = HERE / "data" / "locations.csv"
OUTPUT_DOCX = HERE / "Campus Location Index.docx"

# No. / Name / Digital Address / Code / Category, summing to the 170 mm of
# usable width left by A4 minus 20 mm margins. Word needs the width on every
# cell as well as the column, or it silently re-fits the table.
COL_WIDTHS_MM = [12, 78, 32, 13, 35]
HEADINGS = ["No.", "Location Name", "Digital Address", "Code", "Category"]

HEADER_FILL = "1F3864"   # dark blue
BAND_FILL = "F2F5FA"     # very light blue for alternate rows
PLACEHOLDER = "—"   # em dash, for an address we do not have


def shade(cell, fill):
    """Cell background — python-docx has no API for this, so drop to XML."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def repeat_as_header(row):
    """Re-print this row at the top of every page the table spills onto."""
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(tblHeader)


def add_field(paragraph, code):
    """Insert a Word field (PAGE, NUMPAGES) that updates on open."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, end):
        run._r.append(node)
    return run


def write_cell(cell, text, *, bold=False, align=None, color=None, size=9):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def build(locations, addresses):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)      # A4 — python-docx defaults to US Letter
    section.page_height = Mm(297)
    for side in ("top", "bottom", "left", "right"):
        setattr(section, f"{side}_margin", Mm(20))

    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10)

    title = doc.add_heading("Campus Location Index", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        "University of Ghana, Legon — Campus Service & Routing Database"
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub_run.italic = True

    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run("Format. ").bold = True
    intro.add_run(
        "Each location is indexed as a four-field record of the form "
    )
    fmt = intro.add_run("('New N Block', 'G4-522-1894', '02', 'Lecture Hall')")
    fmt.font.name = "Consolas"
    fmt.font.size = Pt(9)
    intro.add_run(
        " — location name, GhanaPost GPS digital address, two-digit "
        "category code, and category label. The category code is the position "
        "of the category in the schema's permitted-category list, so the codes "
        "shift if that list is reordered."
    )

    note = doc.add_paragraph()
    note.add_run("Digital addresses. ").bold = True
    note.add_run(
        "The source dataset records latitude and longitude only, so no "
        "GhanaPost GPS addresses are available yet. Those cells are shown as "
        "— and are to be filled once the addresses are collected; they "
        "have deliberately not been estimated."
    )

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(HEADINGS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for column, width_mm in zip(table.columns, COL_WIDTHS_MM):
        column.width = Mm(width_mm)

    header = table.rows[0]
    repeat_as_header(header)
    for cell, heading in zip(header.cells, HEADINGS):
        shade(cell, HEADER_FILL)
        write_cell(
            cell,
            heading,
            bold=True,
            color=RGBColor(0xFF, 0xFF, 0xFF),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    centre = WD_ALIGN_PARAGRAPH.CENTER
    for number, loc in enumerate(locations, start=1):
        category = loc["category"]
        row = table.add_row()
        values = [
            (str(number), centre),
            (loc["name"], None),
            (addresses.get(loc["location_code"]) or PLACEHOLDER, centre),
            (CODE[category], centre),
            (LABEL[category], None),
        ]
        for cell, width_mm, (text, align) in zip(row.cells, COL_WIDTHS_MM, values):
            cell.width = Mm(width_mm)   # required alongside the column width
            write_cell(cell, text, align=align)
            if number % 2 == 0:
                shade(cell, BAND_FILL)

    doc.add_paragraph()
    total = doc.add_paragraph()
    total_run = total.add_run(f"Total: {len(locations)} indexed locations.")
    total_run.bold = True
    total_run.font.size = Pt(10)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE")
    footer.add_run(" of ")
    add_field(footer, "NUMPAGES")
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    return doc


def main():
    addresses = load_addresses(sys.argv[1] if len(sys.argv) > 1 else None)

    with open(LOCATIONS_CSV, newline="", encoding="utf-8") as fh:
        locations = [row for row in csv.DictReader(fh) if row["category"] in CODE]

    build(locations, addresses).save(OUTPUT_DOCX)

    filled = sum(1 for loc in locations if addresses.get(loc["location_code"]))
    print(f"Wrote {OUTPUT_DOCX.name}: {len(locations)} locations, "
          f"{filled} digital addresses filled")


if __name__ == "__main__":
    main()
